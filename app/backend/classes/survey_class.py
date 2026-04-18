from app.backend.db.models import SurveyModel, SurveyQuestionModel, SurveyQuestionOptionModel, SurveyResponseAnswerModel
from datetime import datetime
import json
import os
from pathlib import Path

class SurveyClass:
    def __init__(self, db):
        self.db = db

    # ========== SURVEYS ==========
    
    def get_all_surveys(self):
        """Obtiene todas las encuestas"""
        try:
            surveys = (
                self.db.query(SurveyModel)
                .filter(SurveyModel.status_id == 1)
                .order_by(SurveyModel.created_at.desc())
                .all()
            )
            return [{
                "id": survey.id,
                "branch_office_id": survey.branch_office_id,
                "status_id": survey.status_id,
                "title": survey.title,
                "description": survey.description,
                "created_at": survey.created_at.isoformat() if survey.created_at else None,
                "updated_at": survey.updated_at.isoformat() if survey.updated_at else None
            } for survey in surveys]
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def get_survey(self, survey_id):
        """Obtiene una encuesta con sus preguntas y opciones"""
        try:
            survey = self.db.query(SurveyModel).filter(SurveyModel.id == survey_id).first()
            if not survey:
                return {"status": "error", "message": "Survey not found"}
            
            # Obtener preguntas ordenadas
            questions = self.db.query(SurveyQuestionModel).filter(
                SurveyQuestionModel.survey_id == survey_id
            ).order_by(SurveyQuestionModel.order).all()
            
            questions_data = []
            for question in questions:
                question_dict = {
                    "id": question.id,
                    "question": question.question,
                    "field_type": question.field_type,
                    "order": question.order
                }
                
                # Si la pregunta tiene opciones (select, radio, checkbox), obtenerlas
                if question.field_type in ['select', 'radio', 'checkbox']:
                    options = self.db.query(SurveyQuestionOptionModel).filter(
                        SurveyQuestionOptionModel.question_id == question.id
                    ).order_by(SurveyQuestionOptionModel.order).all()
                    
                    question_dict["options"] = [{
                        "id": option.id,
                        "option_text": option.option_text,
                        "order": option.order
                    } for option in options]
                
                questions_data.append(question_dict)
            
            return {
                "id": survey.id,
                "branch_office_id": survey.branch_office_id,
                "status_id": survey.status_id,
                "title": survey.title,
                "description": survey.description,
                "created_at": survey.created_at.isoformat() if survey.created_at else None,
                "updated_at": survey.updated_at.isoformat() if survey.updated_at else None,
                "questions": questions_data
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def store_survey(self, survey_data):
        """Crea una nueva encuesta con sus preguntas y opciones"""
        try:
            # Crear la encuesta
            survey = SurveyModel(
                branch_office_id=survey_data.get("branch_office_id", 0),
                status_id=survey_data.get("status_id", 1),
                title=survey_data.get("title"),
                description=survey_data.get("description"),
            )
            self.db.add(survey)
            self.db.flush()  # Para obtener el ID
            
            # Crear las preguntas
            questions_data = survey_data.get("questions", [])
            for question_data in questions_data:
                question = SurveyQuestionModel(
                    survey_id=survey.id,
                    question=question_data.get("question"),
                    field_type=question_data.get("field_type", "text"),
                    order=question_data.get("order", 0)
                )
                self.db.add(question)
                self.db.flush()  # Para obtener el ID
                
                # Si la pregunta tiene opciones, crearlas
                if question.field_type in ['select', 'radio', 'checkbox']:
                    options_data = question_data.get("options", [])
                    for option_data in options_data:
                        option = SurveyQuestionOptionModel(
                            question_id=question.id,
                            option_text=option_data.get("option_text"),
                            order=option_data.get("order", 0)
                        )
                        self.db.add(option)
            
            self.db.commit()
            return {"status": "success", "message": "Survey created successfully", "survey_id": survey.id}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    def update_survey(self, survey_id, survey_data):
        """Actualiza una encuesta"""
        try:
            survey = self.db.query(SurveyModel).filter(SurveyModel.id == survey_id).first()
            if not survey:
                return {"status": "error", "message": "Survey not found"}
            
            survey.title = survey_data.get("title", survey.title)
            survey.description = survey_data.get("description", survey.description)
            if "branch_office_id" in survey_data:
                survey.branch_office_id = survey_data["branch_office_id"]
            if "status_id" in survey_data:
                survey.status_id = survey_data["status_id"]
            survey.updated_at = datetime.utcnow()
            
            self.db.commit()
            return {"status": "success", "message": "Survey updated successfully"}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    def delete_survey(self, survey_id):
        """Elimina una encuesta (cascade elimina preguntas, opciones y respuestas)"""
        try:
            survey = self.db.query(SurveyModel).filter(SurveyModel.id == survey_id).first()
            if not survey:
                return {"status": "error", "message": "Survey not found"}
            
            self.db.delete(survey)
            self.db.commit()
            return {"status": "success", "message": "Survey deleted successfully"}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    # ========== QUESTIONS ==========
    
    def add_question(self, survey_id, question_data):
        """Agrega una pregunta a una encuesta"""
        try:
            survey = self.db.query(SurveyModel).filter(SurveyModel.id == survey_id).first()
            if not survey:
                return {"status": "error", "message": "Survey not found"}
            
            question = SurveyQuestionModel(
                survey_id=survey_id,
                question=question_data.get("question"),
                field_type=question_data.get("field_type", "text"),
                order=question_data.get("order", 0)
            )
            self.db.add(question)
            self.db.flush()
            
            # Si la pregunta tiene opciones, crearlas
            if question.field_type in ['select', 'radio', 'checkbox']:
                options_data = question_data.get("options", [])
                for option_data in options_data:
                    option = SurveyQuestionOptionModel(
                        question_id=question.id,
                        option_text=option_data.get("option_text"),
                        order=option_data.get("order", 0)
                    )
                    self.db.add(option)
            
            self.db.commit()
            return {"status": "success", "message": "Question added successfully", "question_id": question.id}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    def update_question(self, question_id, question_data):
        """Actualiza una pregunta"""
        try:
            question = self.db.query(SurveyQuestionModel).filter(SurveyQuestionModel.id == question_id).first()
            if not question:
                return {"status": "error", "message": "Question not found"}
            
            question.question = question_data.get("question", question.question)
            question.field_type = question_data.get("field_type", question.field_type)
            question.order = question_data.get("order", question.order)
            question.updated_at = datetime.utcnow()
            
            # Si cambió el tipo y ahora tiene opciones, crearlas
            if question.field_type in ['select', 'radio', 'checkbox']:
                options_data = question_data.get("options", [])
                # Eliminar opciones existentes
                self.db.query(SurveyQuestionOptionModel).filter(
                    SurveyQuestionOptionModel.question_id == question_id
                ).delete()
                # Crear nuevas opciones
                for option_data in options_data:
                    option = SurveyQuestionOptionModel(
                        question_id=question_id,
                        option_text=option_data.get("option_text"),
                        order=option_data.get("order", 0)
                    )
                    self.db.add(option)
            
            self.db.commit()
            return {"status": "success", "message": "Question updated successfully"}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    def delete_question(self, question_id):
        """Elimina una pregunta (cascade elimina opciones y respuestas)"""
        try:
            question = self.db.query(SurveyQuestionModel).filter(SurveyQuestionModel.id == question_id).first()
            if not question:
                return {"status": "error", "message": "Question not found"}
            
            self.db.delete(question)
            self.db.commit()
            return {"status": "success", "message": "Question deleted successfully"}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    # ========== RESPONSES ==========
    
    def submit_response(self, survey_id, responses_data):
        """Guarda las respuestas de una encuesta"""
        try:
            survey = self.db.query(SurveyModel).filter(SurveyModel.id == survey_id).first()
            if not survey:
                return {"status": "error", "message": "Survey not found"}
            
            # responses_data es una lista de objetos: [{"question_id": 1, "answer_text": "...", "option_id": null}, ...]
            for response_data in responses_data:
                question_id = response_data.get("question_id")
                answer_text = response_data.get("answer_text")
                option_id = response_data.get("option_id")
                
                # Validar que la pregunta pertenezca a la encuesta
                question = self.db.query(SurveyQuestionModel).filter(
                    SurveyQuestionModel.id == question_id,
                    SurveyQuestionModel.survey_id == survey_id
                ).first()
                
                if not question:
                    continue  # Saltar si la pregunta no existe o no pertenece a la encuesta
                
                # Validar que si es opción, el option_id sea válido
                if option_id:
                    option = self.db.query(SurveyQuestionOptionModel).filter(
                        SurveyQuestionOptionModel.id == option_id,
                        SurveyQuestionOptionModel.question_id == question_id
                    ).first()
                    if not option:
                        continue  # Saltar si la opción no existe o no pertenece a la pregunta
                
                # Crear la respuesta
                answer = SurveyResponseAnswerModel(
                    survey_id=survey_id,
                    question_id=question_id,
                    answer_text=answer_text,
                    option_id=option_id
                )
                self.db.add(answer)
            
            self.db.commit()
            return {"status": "success", "message": "Response submitted successfully"}
        except Exception as e:
            self.db.rollback()
            return {"status": "error", "message": str(e)}
    
    def get_survey_responses(self, survey_id):
        """Obtiene todas las respuestas de una encuesta"""
        try:
            survey = self.db.query(SurveyModel).filter(SurveyModel.id == survey_id).first()
            if not survey:
                return {"status": "error", "message": "Survey not found"}
            
            responses = self.db.query(SurveyResponseAnswerModel).filter(
                SurveyResponseAnswerModel.survey_id == survey_id
            ).order_by(SurveyResponseAnswerModel.created_at).all()
            
            return [{
                "id": response.id,
                "question_id": response.question_id,
                "answer_text": response.answer_text,
                "option_id": response.option_id,
                "created_at": response.created_at.isoformat() if response.created_at else None
            } for response in responses]
        except Exception as e:
            return {"status": "error", "message": str(e)}

    def export_survey_results_docx(self, survey_id: int):
        """
        Genera un documento Word (.docx) con logo Jisparking, portada, sección
        Introducción (texto fijo breve + descripción de la encuesta si existe),
        y una página por pregunta con gráficos agregados (sin detalle por respuesta).
        Retorna (bytes | None, error_message | None).
        """
        # Colores marca Jisparking (logo / identidad)
        _JIS_BLUE = "#002080"
        _JIS_ORANGE = "#EF7D00"
        _CHART_TEXT_INNER = "#FFFFFF"  # texto sobre barras / porciones de color

        try:
            from io import BytesIO
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        except ImportError:
            return None, "python-docx no está instalado. Ejecute: pip install python-docx"

        try:
            import matplotlib

            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
        except ImportError:
            return None, "matplotlib no está instalado. Ejecute: pip install matplotlib"

        def _resolve_logo_path() -> Path | None:
            base = Path(__file__).resolve().parent.parent / "static"
            for name in ("logo.png", "intrajis_logo.png", "jisparking_logo.png"):
                p = base / name
                if p.is_file():
                    return p
            envp = os.environ.get("SURVEY_DOCX_LOGO_PATH", "").strip()
            if envp:
                pe = Path(envp)
                if pe.is_file():
                    return pe
            return None

        def _apply_branding(document) -> None:
            normal = document.styles["Normal"]
            nf = normal.font
            nf.name = "Calibri"
            nf.size = Pt(11)
            try:
                nf.color.rgb = RGBColor(0x1D, 0x26, 0x30)
            except Exception:
                pass
            for level in (1, 2, 3):
                try:
                    hs = document.styles[f"Heading {level}"]
                    hs.font.name = "Calibri"
                    hs.font.bold = True
                    # Preguntas (nivel 3) en negro; títulos de sección en azul marca
                    if level == 3:
                        hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    else:
                        hs.font.color.rgb = RGBColor(0x00, 0x20, 0x80)
                    hs.font.size = Pt(20 if level == 1 else (14 if level == 2 else 12))
                except KeyError:
                    pass

        def _respondent_count(question: dict, q_answers: list) -> int:
            if question.get("field_type") == "checkbox":
                buckets = []
                for a in q_answers:
                    t = a.get("created_at")
                    if not t:
                        buckets.append(f"row:{a.get('id')}")
                    else:
                        buckets.append(t[:19] if len(t) >= 19 else t)
                return len(set(buckets))
            return len(q_answers)

        def _survey_chart_png(
            question: dict,
            q_answers: list,
            global_option_labels: dict,
        ):
            """Devuelve BytesIO PNG o None si no aplica / sin datos."""
            ft = question.get("field_type") or "text"
            if ft == "text":
                return None

            option_counts: dict = {}
            for a in q_answers:
                oid = a.get("option_id")
                if oid:
                    option_counts[oid] = option_counts.get(oid, 0) + 1

            total_resp = _respondent_count(question, q_answers)
            if total_resp == 0 or not option_counts:
                return None

            labels: list = []
            counts: list = []
            opts = question.get("options") or []
            if opts:
                for opt in sorted(opts, key=lambda x: (x.get("order", 0), x.get("id") or 0)):
                    oid = opt.get("id")
                    labels.append(str(opt.get("option_text") or f"#{oid}"))
                    counts.append(option_counts.get(oid, 0))
            else:
                for oid in sorted(option_counts.keys()):
                    labels.append(
                        str(global_option_labels.get(oid) or f"Opción #{oid}")
                    )
                    counts.append(option_counts[oid])

            if not counts or sum(counts) == 0:
                return None

            percentages = [
                (c / total_resp) * 100.0 if total_resp else 0.0 for c in counts
            ]

            fig_w = min(10, max(6, 0.45 * len(labels) + 3))
            fig, ax = plt.subplots(figsize=(fig_w, 4.2), dpi=120)
            pie_palette = [_JIS_BLUE, _JIS_ORANGE, "#5B6B79", "#2ca87f", "#e58a00", "#3ec9d6"]

            if ft == "radio":
                colors = [pie_palette[i % len(pie_palette)] for i in range(len(labels))]
                _w, texts, autotexts = ax.pie(
                    counts,
                    labels=labels,
                    autopct=lambda pct: f"{pct:.1f}%",
                    startangle=90,
                    colors=colors,
                    textprops={"fontsize": 8, "color": _JIS_BLUE},
                )
                ax.axis("equal")
                for t in texts:
                    t.set_color(_JIS_BLUE)
                for t in autotexts:
                    t.set_color(_CHART_TEXT_INNER)
                    t.set_fontweight("bold")
                ax.set_title(
                    "Distribución de respuestas (torta)",
                    fontsize=11,
                    color=_JIS_BLUE,
                    fontweight="bold",
                )
            else:
                x_pos = range(len(labels))
                bars = ax.bar(
                    x_pos,
                    percentages,
                    color=_JIS_ORANGE,
                    edgecolor=_JIS_BLUE,
                    linewidth=1.0,
                )
                ax.set_xticks(list(x_pos))
                ax.set_xticklabels(labels, rotation=28, ha="right", fontsize=8, color=_JIS_BLUE)
                ax.set_ylabel("Porcentaje (%)", fontsize=9, color=_JIS_BLUE)
                ax.set_ylim(0, 105)
                ax.tick_params(axis="y", colors=_JIS_BLUE)
                ax.set_title(
                    "Distribución de respuestas (barras — % sobre encuestados)",
                    fontsize=11,
                    color=_JIS_BLUE,
                    fontweight="bold",
                )
                for bar, pct, cnt in zip(bars, percentages, counts):
                    h = bar.get_height()
                    cx = bar.get_x() + bar.get_width() / 2.0
                    label = f"{pct:.1f}%\n(n={cnt})"
                    if h >= 10:
                        ax.text(
                            cx,
                            h / 2.0,
                            label,
                            ha="center",
                            va="center",
                            fontsize=7,
                            color=_CHART_TEXT_INNER,
                            fontweight="bold",
                        )
                    else:
                        ax.text(
                            cx,
                            min(h + 2, 98),
                            label,
                            ha="center",
                            va="bottom",
                            fontsize=7,
                            color=_JIS_BLUE,
                        )

            plt.tight_layout()
            out = BytesIO()
            fig.savefig(out, format="png", bbox_inches="tight", facecolor="white")
            plt.close(fig)
            out.seek(0)
            return out

        try:
            survey = self.get_survey(survey_id)
            if isinstance(survey, dict) and survey.get("status") == "error":
                return None, survey.get("message", "Encuesta no encontrada")
            if not isinstance(survey, dict) or survey.get("id") is None:
                return None, "Encuesta no encontrada"

            raw_resp = self.get_survey_responses(survey_id)
            if isinstance(raw_resp, dict) and raw_resp.get("status") == "error":
                return None, raw_resp.get("message", "Error al leer respuestas")
            responses = raw_resp if isinstance(raw_resp, list) else []

            option_map = {}
            for q in survey.get("questions") or []:
                for opt in q.get("options") or []:
                    oid = opt.get("id")
                    if oid is not None:
                        option_map[oid] = opt.get("option_text") or ""

            questions = sorted(
                survey.get("questions") or [],
                key=lambda x: (x.get("order") if x.get("order") is not None else 0, x.get("id") or 0),
            )

            doc = Document()
            _apply_branding(doc)
            # python-docx 1.1.x puede dejar el cuerpo sin <w:p> (solo sectPr): paragraphs[] vacío.
            # Versiones anteriores traían un párrafo vacío inicial. No indexar [0] sin comprobar.
            if len(doc.paragraphs) > 0:
                p0 = doc.paragraphs[0]
                p0._element.getparent().remove(p0._element)

            def _meses_es() -> tuple:
                return (
                    "enero",
                    "febrero",
                    "marzo",
                    "abril",
                    "mayo",
                    "junio",
                    "julio",
                    "agosto",
                    "septiembre",
                    "octubre",
                    "noviembre",
                    "diciembre",
                )

            def _fecha_larga_es(iso_dt: str | None) -> str:
                meses = _meses_es()
                if iso_dt:
                    try:
                        s = str(iso_dt).replace("Z", "+00:00")
                        d = datetime.fromisoformat(s)
                        return f"{d.day} de {meses[d.month - 1]} de {d.year}"
                    except Exception:
                        pass
                from datetime import date

                d = date.today()
                return f"{d.day} de {meses[d.month - 1]} de {d.year}"

            def _cover_spacer(docu, after_pt: int = 10) -> None:
                p = docu.add_paragraph()
                p.paragraph_format.space_after = Pt(after_pt)

            def _cover_center_run(
                docu,
                text: str,
                *,
                size_pt: int,
                bold: bool = False,
                italic: bool = False,
                color_rgb: tuple[int, int, int] | None = None,
            ) -> None:
                p = docu.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.bold = bold
                r.italic = italic
                r.font.size = Pt(size_pt)
                r.font.name = "Calibri"
                if color_rgb:
                    r.font.color.rgb = RGBColor(*color_rgb)

            # ----- Portada
            _cover_spacer(doc, 36)
            logo_path = _resolve_logo_path()
            if logo_path:
                pl = doc.add_paragraph()
                pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pl.paragraph_format.space_after = Pt(8)
                pl.add_run().add_picture(str(logo_path), width=Inches(1.85))

            _cover_spacer(doc, 18)
            stitle = str(survey.get("title") or "Encuesta")
            p_main = doc.add_paragraph()
            p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_main.paragraph_format.space_after = Pt(6)
            rm = p_main.add_run(stitle)
            rm.bold = True
            rm.font.size = Pt(26)
            rm.font.name = "Calibri"
            rm.font.color.rgb = RGBColor(0x00, 0x20, 0x80)

            _cover_center_run(
                doc,
                "Informe de resultados de encuesta",
                size_pt=14,
                italic=True,
                color_rgb=(0, 32, 128),
            )
            _cover_spacer(doc, 6)

            p_rule = doc.add_paragraph()
            p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_rule.paragraph_format.space_after = Pt(14)
            rr = p_rule.add_run("― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ―")
            rr.font.size = Pt(11)
            rr.font.color.rgb = RGBColor(0xEF, 0x7D, 0x00)

            ref_dt = survey.get("updated_at") or survey.get("created_at")
            _cover_center_run(
                doc,
                f"Generado el {_fecha_larga_es(ref_dt)}",
                size_pt=11,
                color_rgb=(29, 38, 48),
            )
            sid = survey.get("id")
            if sid is not None:
                _cover_center_run(
                    doc,
                    f"Identificador de encuesta · #{sid}",
                    size_pt=10,
                    color_rgb=(91, 107, 121),
                )

            _cover_spacer(doc, 28)
            _cover_center_run(
                doc,
                "IntraJIS · Jisparking",
                size_pt=9,
                color_rgb=(91, 107, 121),
            )
            _cover_center_run(
                doc,
                "Documento generado de forma electrónica",
                size_pt=8,
                italic=True,
                color_rgb=(150, 160, 170),
            )

            # Salto a cuerpo del informe
            p_br = doc.add_paragraph()
            p_br.paragraph_format.space_before = Pt(24)
            br = p_br.add_run()
            br.add_break(WD_BREAK.PAGE)

            doc.add_heading(str(survey.get("title") or "Encuesta"), level=1)

            doc.add_heading("Introducción", level=2)

            def _justify_par(docu, text: str, space_after_pt: int = 11) -> None:
                pr = docu.add_paragraph()
                pr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pr.paragraph_format.space_after = Pt(space_after_pt)
                pr.paragraph_format.line_spacing = 1.15
                pr.add_run(text)

            _justify_par(
                doc,
                "El presente documento constituye un informe ejecutivo elaborado a partir de "
                "las respuestas registradas en la encuesta. Su finalidad es entregar una visión "
                "ordenada y accesible de los resultados, de modo que las personas responsables del "
                "servicio o de la operación puedan identificar con rapidez tendencias, prioridades "
                "de mejora y oportunidades de acción alineadas con la experiencia del usuario.",
            )
            _justify_par(
                doc,
                "La información se presenta de manera sintética y orientada a la toma de decisiones: "
                "para las preguntas con alternativas de respuesta se incorporan gráficos que "
                "resumen la distribución de las opciones elegidas; en el caso de preguntas de "
                "texto libre se indica el volumen de aportes recibidos, sin reproducir el detalle "
                "literal de cada participante. Cada pregunta se desarrolla en una página distinta, "
                "lo que facilita la lectura, la impresión por secciones y el archivo del informe.",
            )
            _justify_par(
                doc,
                "Los resultados deben interpretarse en el contexto operativo y temporal en que se "
                "aplicó el instrumento; no sustituyen, por sí mismos, un análisis cualitativo "
                "complementario cuando la naturaleza del tema lo requiera. Se sugiere utilizar este "
                "informe como insumo en reuniones de seguimiento, planes de mejora continua y en la "
                "comunicación interna o con partes interesadas, preservando la confidencialidad que "
                "corresponda a los datos tratados.",
            )

            desc = survey.get("description")
            if desc and str(desc).strip():
                p_ctx = doc.add_paragraph()
                p_ctx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_ctx.paragraph_format.space_before = Pt(6)
                p_ctx.paragraph_format.space_after = Pt(12)
                p_ctx.paragraph_format.line_spacing = 1.15
                r_ctx = p_ctx.add_run("Contexto de la encuesta: ")
                r_ctx.bold = True
                r_ctx.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                p_ctx.add_run(str(desc).strip())

            doc.add_paragraph("")

            for num, q in enumerate(questions, start=1):
                # Una pregunta (y su gráfico / texto) por página
                p_new = doc.add_paragraph()
                r_new = p_new.add_run()
                r_new.add_break(WD_BREAK.PAGE)

                qid = q.get("id")
                qtext = q.get("question") or f"Pregunta {qid}"
                heading_para = doc.add_heading(f"{num}. {qtext}", level=3)
                for _run in heading_para.runs:
                    _run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                heading_para.paragraph_format.space_after = Pt(6)

                q_answers = [r for r in responses if r.get("question_id") == qid]
                if not q_answers:
                    doc.add_paragraph("(Sin respuestas)")
                    heading_para.paragraph_format.space_after = Pt(14)
                    continue

                ft = q.get("field_type") or "text"
                if ft == "text":
                    n_txt = len(q_answers)
                    doc.add_paragraph(
                        f"Esta pregunta admite texto libre. Respuestas registradas: {n_txt} "
                        "(el detalle no se incluye en este informe)."
                    )
                    heading_para.paragraph_format.space_after = Pt(14)
                    continue

                chart_buf = _survey_chart_png(q, q_answers, option_map)
                if chart_buf:
                    gap = doc.add_paragraph()
                    gap.paragraph_format.space_before = Pt(28)
                    gap.paragraph_format.space_after = Pt(10)
                    doc.add_picture(chart_buf, width=Inches(5.9))
                    after_pic = doc.paragraphs[-1]
                    after_pic.paragraph_format.space_after = Pt(22)
                else:
                    p = doc.add_paragraph()
                    r = p.add_run("(Sin datos suficientes para generar el gráfico.)")
                    r.italic = True
                    heading_para.paragraph_format.space_after = Pt(14)

            bio = BytesIO()
            doc.save(bio)
            return bio.getvalue(), None
        except Exception as e:
            return None, str(e)
